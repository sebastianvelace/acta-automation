from __future__ import annotations

from pathlib import Path

from docx import Document

from src.parser import _extract_times_from_filename, extract_text


def test_extract_times_from_filename_space_format() -> None:
    hora, fin = _extract_times_from_filename(
        "Actualización Dashboard - Universal _ 2026_05_21 11_02 GMT-05_00 - Notas de Gemini.docx"
    )
    assert hora == "11:02 AM"
    assert fin == ""


def test_extract_text_populates_gorila_teams_and_emails(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph(
        "Invitado camilolinaresj@gmail.com Marketing Gorila Hosting "
        "Administración Gorila Hosting ads@gorilahosting@gmail.com Social Media Gorila Hosting"
    )
    doc.add_paragraph("Archivos adjuntos x")
    p = tmp_path / "invite.docx"
    doc.save(str(p))
    out = extract_text(str(p))
    meta = out["metadata"]
    assert "camilolinaresj@gmail.com" in meta["attendee_emails"]
    assert "camilolinaresj@gmail.com" in meta["client_emails"]
    assert "gorilahosting@gmail.com" in meta["attendee_emails"]
    assert "gorilahosting@gmail.com" in meta["gorila_emails"]
    assert "Marketing Gorila Hosting" in meta["gorila_teams"]
    assert "Administración Gorila Hosting" in meta["gorila_teams"]


def test_detect_virtual_from_grabacion(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Registros de la reunión Grabación")
    p = tmp_path / "virtual.docx"
    doc.save(str(p))
    out = extract_text(str(p))
    assert out["metadata"]["is_virtual"] is True


def test_skips_gemini_export_timestamp_as_hora_inicio(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("may 25, 2026")
    doc.add_paragraph("Invitado pedro@example.com")
    for minute in ("59", "58"):
        name = (
            f"Reunión Reporte de Ventas - Universal _ "
            f"2026_05_25 15_{minute} GMT-05_00 - Notas de Gemini.docx"
        )
        p = tmp_path / name
        doc.save(str(p))
        out = extract_text(str(p))
        if minute == "59":
            assert out["metadata"]["hora_inicio"] == ""
        else:
            assert out["metadata"]["hora_inicio"] == "3:58 PM"


def test_keeps_meeting_time_from_filename_when_not_export_stamp(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("may 21, 2026")
    name = (
        "Actualización Dashboard - Universal _ "
        "2026_05_21 11_02 GMT-05_00 - Notas de Gemini.docx"
    )
    p = tmp_path / name
    doc.save(str(p))
    out = extract_text(str(p))
    assert out["metadata"]["hora_inicio"] == "11:02 AM"
