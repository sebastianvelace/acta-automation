from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
from docx import Document

from src.aliases import post_process_acta
from src.pipeline import run_acta_pipeline
from src.schemas import ActaSchema

FIXTURES = Path(__file__).resolve().parent / "fixtures"
GEMINI = FIXTURES / "gemini"
EXPECTED = FIXTURES / "expected"


def _gather_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            chunks.append(p.text)
        for table in sec.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        for p in sec.footer.paragraphs:
            chunks.append(p.text)
    return "\n".join(chunks)


def _norm_snippet(s: str) -> str:
    s = s.replace("&amp;", " ").replace("&", " ")
    return re.sub(r"\s+", " ", s).strip().casefold()


def _docx_contains(docx_path: Path, needle: str) -> bool:
    hay = _gather_docx_text(docx_path)
    return _norm_snippet(needle) in _norm_snippet(hay)


def _structure_meeting_like_production(payload: dict) -> dict:
    return post_process_acta(ActaSchema.model_validate(payload).model_dump())


@pytest.fixture
def mock_llm(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_structure_meeting(
        raw_text: str,
        metadata: dict | None = None,
        source_filename: str | None = None,
    ) -> dict:
        stem = Path(source_filename or "x").stem
        path = EXPECTED / f"{stem}.llm_raw.json"
        data = json.loads(path.read_text(encoding="utf-8"))
        return _structure_meeting_like_production(data)

    monkeypatch.setattr("src.pipeline.structure_meeting", fake_structure_meeting)


@pytest.mark.parametrize("stem", [f"case_{i:02d}" for i in range(1, 6)])
def test_run_acta_pipeline_generates_outputs(
    stem: str,
    mock_llm: None,
    tmp_path: Path,
) -> None:
    docx_in = GEMINI / f"{stem}.docx"
    assert docx_in.is_file(), f"Missing fixture {docx_in}"

    raw_payload = json.loads((EXPECTED / f"{stem}.llm_raw.json").read_text(encoding="utf-8"))
    expected_cliente = raw_payload["cliente"]

    result = run_acta_pipeline(
        str(docx_in),
        source_filename=f"{stem}.docx",
        output_dir=str(tmp_path),
        keep_docx=True,
    )

    pdf_path = Path(result["pdf_path"])
    docx_out = Path(result["docx_path"] or "")

    assert pdf_path.is_file()
    assert pdf_path.stat().st_size > 10 * 1024
    assert docx_out.is_file()

    assert _docx_contains(docx_out, expected_cliente)
