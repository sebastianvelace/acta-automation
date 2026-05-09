#!/usr/bin/env python3
"""Genera tests/fixtures/gemini/*.docx y deja listo el árbol expected (JSON editados aparte si hace falta)."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
GEMINI = ROOT / "tests/fixtures/gemini"
PROCESSED = (
    ROOT
    / "input/processed/Seguimiento - Eventos & Matrimonios_ 2026_05_06 16_07 GMT-05_00 - Notas de Gemini.docx"
)


def _gemini_paragraphs(
    date_line: str,
    title: str,
    invitados_csv: str,
    *,
    extra_body: list[str] | None = None,
) -> Document:
    d = Document()
    d.add_paragraph(date_line)
    d.add_paragraph(title)
    d.add_paragraph(f"Invitados: {invitados_csv}")
    d.add_paragraph("Archivos adjuntos")
    d.add_paragraph("Detalles")
    for line in extra_body or [
        "Hora de inicio: 10:00 AM",
        "Reunión interna de seguimiento.",
    ]:
        d.add_paragraph(line)
    return d


def main() -> None:
    GEMINI.mkdir(parents=True, exist_ok=True)
    if PROCESSED.is_file():
        shutil.copy(PROCESSED, GEMINI / "case_01.docx")
    else:
        _gemini_paragraphs(
            "may 6, 2026",
            "Seguimiento - Eventos & Matrimonios",
            "Social Media Gorila Hosting, Marketing Gorila Hosting",
        ).save(GEMINI / "case_01.docx")

    specs = [
        (
            "case_02.docx",
            "jun 1, 2026",
            "Administración - Cierre mensual",
            "Administración Gorila Hosting, Redes Gorila Hosting, cliente@empresa.co",
        ),
        (
            "case_03.docx",
            "jul 15, 2026",
            "Executive review",
            "Executive Gorila Hosting, Soporte Gorila Hosting, Ventas Gorila Hosting",
        ),
        (
            "case_04.docx",
            "ago 3, 2026",
            "Diseño y Producto",
            "Diseño Gorila Hosting, Producto Gorila Hosting, Gorila Hosting",
        ),
        (
            "case_05.docx",
            "sep 9, 2026",
            "Portal cliente",
            "Eventos y Matrominios Portal, Eventos & Matrimonios Portal, ana.garcia@cliente.com",
        ),
    ]
    for filename, date_l, title, inv in specs:
        _gemini_paragraphs(date_l, title, inv).save(GEMINI / filename)


if __name__ == "__main__":
    main()
