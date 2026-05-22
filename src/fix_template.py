import re
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "templates", "acta_template.docx"
)
JINJA_PATTERN = re.compile(r'{{|}}|{%|%}')


def _row_cells_joined(row) -> str:
    return "".join(
        (c.paragraphs[0].text if c.paragraphs else "") for c in row.cells
    )


def _row_has_combined_tr_loop(row) -> bool:
    t = _row_cells_joined(row)
    return "{%tr for " in t and "{%tr endfor %" in t


def fix_paragraph(paragraph):
    runs = paragraph.runs
    if not runs:
        return None
    full_text = "".join(r.text for r in runs)
    if not JINJA_PATTERN.search(full_text):
        return None
    font_name = runs[0].font.name
    font_size = runs[0].font.size
    font_bold = runs[0].font.bold
    for run in runs:
        run.text = ""
    runs[0].text = full_text
    if font_name:
        runs[0].font.name = font_name
    if font_size:
        runs[0].font.size = font_size
    if font_bold is not None:
        runs[0].font.bold = font_bold
    return full_text


def set_cell_text(cell, text):
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def _packed_data_cell_texts(
    texts: list[str], for_cell: int, end_cell: int, ncells: int
) -> list[str]:
    """Place template variables in columns 0..k-1 so they align with table headers (no leading gap)."""
    body = [texts[i] for i in range(ncells) if i not in (for_cell, end_cell)]
    while len(body) < ncells:
        body.append("")
    return body[:ncells]


def split_tr_loop_row(table, data_row_index):
    """
    docxtpl forbids two {%tr ... %} tags in the same table row. One row must be
    only {%tr for ... %}, the next row(s) the body, then a row with only {%tr endfor %}.
    """
    from copy import deepcopy

    row = table.rows[data_row_index]
    ncells = len(row.cells)
    texts = [c.paragraphs[0].text if c.paragraphs else "" for c in row.cells]

    for_cell = next(i for i, t in enumerate(texts) if "{%tr for " in t)
    end_cell = next(i for i, t in enumerate(texts) if "{%tr endfor %" in t)

    for_tag = texts[for_cell].strip()
    end_tag = texts[end_cell].strip()
    packed = _packed_data_cell_texts(texts, for_cell, end_cell, ncells)

    tr_data = row._tr
    tr_open = deepcopy(tr_data)
    tr_data.addprevious(tr_open)

    open_row = table.rows[data_row_index]
    data_row = table.rows[data_row_index + 1]
    for i, c in enumerate(open_row.cells):
        set_cell_text(c, for_tag if i == for_cell else "")
    for i, c in enumerate(data_row.cells):
        set_cell_text(c, packed[i])

    tr_close = deepcopy(data_row._tr)
    data_row._tr.addnext(tr_close)
    close_row = table.rows[data_row_index + 2]
    for i, c in enumerate(close_row.cells):
        set_cell_text(c, end_tag if i == end_cell else "")


def add_column(table, width=800):
    tbl = table._tbl
    tbl_grid = tbl.find(qn('w:tblGrid'))
    grid_col = OxmlElement('w:gridCol')
    grid_col.set(qn('w:w'), str(width))
    tbl_grid.append(grid_col)
    for row in table.rows:
        tc = OxmlElement('w:tc')
        tcp = OxmlElement('w:tcPr')
        tcw = OxmlElement('w:tcW')
        tcw.set(qn('w:w'), str(width))
        tcw.set(qn('w:type'), 'dxa')
        tcp.append(tcw)
        tc.append(tcp)
        tc.append(OxmlElement('w:p'))
        row._tr.append(tc)


def fix_document(path):
    doc = Document(path)

    # Step 0: body clearance under corporate header (avoid PDF overlap with banner)
    print("-- Step 0: Page layout (margins / header gap / title spacing) --")
    for sec in doc.sections:
        sec.top_margin = Pt(102)
        sec.header_distance = Pt(42)
        print(f"  top_margin={sec.top_margin.pt:.0f}pt, header_distance={sec.header_distance.pt:.0f}pt")
    for para in doc.paragraphs:
        raw = para.text or ""
        if "ACTA DE REUNION" in raw.upper():
            para.paragraph_format.space_before = Pt(14)
            print("  space_before ACTA title: 14pt")
            break

    # Step 1: merge split Jinja2 runs
    print("\n-- Step 1: Merging split runs --")
    for para in doc.paragraphs:
        result = fix_paragraph(para)
        if result:
            print(f"  paragraph: {result!r}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    result = fix_paragraph(para)
                    if result:
                        print(f"  cell: {result!r}")

    # Step 2: fix table structure
    print("\n-- Step 2: Fixing table structure --")

    # Table 1 — invitados (3 cols → 4)
    t_invitados = doc.tables[1]
    if len(t_invitados.columns) == 3:
        add_column(t_invitados, width=800)
        print("  invitados: added column 4")
    if len(t_invitados.rows) == 2:
        row = t_invitados.rows[1]
        set_cell_text(row.cells[0], "{%tr for a in invitados %}")
        set_cell_text(row.cells[1], "{{a.nombre}}")
        set_cell_text(row.cells[2], "{{a.puesto}}")
        set_cell_text(row.cells[3], "{%tr endfor %}")
        print("  invitados row 1:", [c.paragraphs[0].text for c in row.cells])
    else:
        print("  invitados: skip row 1 layout (expected 2 rows, have %s)" % len(t_invitados.rows))

    # Table 2 — compromisos_gorila (4 cols → 5)
    t_gorila = doc.tables[2]
    if len(t_gorila.columns) == 4:
        add_column(t_gorila, width=800)
        print("  compromisos_gorila: added column 5")
    if len(t_gorila.rows) == 2:
        row = t_gorila.rows[1]
        set_cell_text(row.cells[0], "{%tr for c in compromisos_gorila %}")
        set_cell_text(row.cells[1], "{{loop.index}}")
        set_cell_text(row.cells[2], "{{c.tarea}}")
        set_cell_text(row.cells[3], "{{c.responsable}}")
        set_cell_text(row.cells[4], "{%tr endfor %}")
        print("  compromisos_gorila row 1:", [c.paragraphs[0].text for c in row.cells])
    else:
        print("  compromisos_gorila: skip row 1 layout (have %s rows)" % len(t_gorila.rows))

    # Table 3 — compromisos_cliente (4 cols → 5)
    t_cliente = doc.tables[3]
    if len(t_cliente.columns) == 4:
        add_column(t_cliente, width=800)
        print("  compromisos_cliente: added column 5")
    if len(t_cliente.rows) == 2:
        row = t_cliente.rows[1]
        set_cell_text(row.cells[0], "{%tr for c in compromisos_cliente %}")
        set_cell_text(row.cells[1], "{{loop.index}}")
        set_cell_text(row.cells[2], "{{c.tarea}}")
        set_cell_text(row.cells[3], "{{c.responsable}}")
        set_cell_text(row.cells[4], "{%tr endfor %}")
        print("  compromisos_cliente row 1:", [c.paragraphs[0].text for c in row.cells])
    else:
        print("  compromisos_cliente: skip row 1 layout (have %s rows)" % len(t_cliente.rows))

    print("\n-- Step 2b: Split {%tr for%}/{%tr endfor%} into separate rows (docxtpl) --")
    for label, idx in (
        ("invitados", 1),
        ("compromisos_gorila", 2),
        ("compromisos_cliente", 3),
    ):
        tbl = doc.tables[idx]
        if len(tbl.rows) > 1 and _row_has_combined_tr_loop(tbl.rows[1]):
            split_tr_loop_row(tbl, 1)
            print(f"  {label}: split row loop at template row index 1")
        else:
            print(f"  {label}: skip split (already split or no combined tr row)")

    print("\n-- Step 2c: Align data rows with headers (pack loop body into cols 0..) --")

    def pack_table_data_row(tbl, label: str) -> None:
        if len(tbl.rows) != 4:
            return
        row = tbl.rows[2]
        texts = [c.paragraphs[0].text if c.paragraphs else "" for c in row.cells]
        if label == "invitados":
            cleaned = []
            for t in texts:
                s = (t or "").replace("✓", "").strip()
                cleaned.append(s)
            texts = cleaned
        body = [t for t in texts if t.strip()]
        for i, c in enumerate(row.cells):
            set_cell_text(c, body[i] if i < len(body) else "")
        if label == "invitados" and len(row.cells) > 2:
            set_cell_text(row.cells[2], "Confirmado")
        print(f"  {label}: data row packed ({len(body)} fields)")

    pack_table_data_row(doc.tables[1], "invitados")
    pack_table_data_row(doc.tables[2], "compromisos_gorila")
    pack_table_data_row(doc.tables[3], "compromisos_cliente")

    # Step 3: asuntos_tratados — {%p for%} / titulo (bold) / descripcion / {%p endfor%}
    print("\n-- Step 3: Splitting asuntos_tratados into four paragraphs --")
    TARGET = "{%p for a in asuntos_tratados %}"
    TITULO_TEMPLATE = "{{loop.index}}. {{a.titulo}}"
    DESCR_TEMPLATE = "{{a.descripcion}}"
    LEGACY_CONTENT = "{{loop.index}}. {{a.titulo}}: {{a.descripcion}}"
    ENDFOR = "{%p endfor %}"

    def _append_plain_jinja_run(paragraph_el, text):
        r_el = OxmlElement("w:r")
        t_el = OxmlElement("w:t")
        t_el.text = text
        r_el.append(t_el)
        paragraph_el.append(r_el)

    def _insert_jinja_paragraph(
        parent, insert_idx, style_name, text, *, space_before_pt=None
    ):
        p_el = OxmlElement("w:p")
        pPr_elems = []
        if style_name:
            pStyle = OxmlElement("w:pStyle")
            pStyle.set(qn("w:val"), style_name)
            pPr_elems.append(pStyle)
        if space_before_pt is not None:
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
            pPr_elems.append(spacing)
        if pPr_elems:
            pPr = OxmlElement("w:pPr")
            for el in pPr_elems:
                pPr.append(el)
            p_el.append(pPr)
        _append_plain_jinja_run(p_el, text)
        parent.insert(insert_idx, p_el)

    def _apply_titulo_bold_and_spacing(titulo_paragraph):
        for run in titulo_paragraph.runs:
            run.text = ""
        r = titulo_paragraph.add_run(TITULO_TEMPLATE)
        r.bold = True
        titulo_paragraph.paragraph_format.space_before = Pt(12)

    def _bind_titulo_paragraph():
        for p in doc.paragraphs:
            if TITULO_TEMPLATE in p.text:
                _apply_titulo_bold_and_spacing(p)
                return True
        return False

    found = False
    paras = list(doc.paragraphs)

    # Case A — already four paragraphs (for / titulo / descr / endfor)
    for i in range(len(paras) - 3):
        t0 = "".join(r.text for r in paras[i].runs)
        t1 = "".join(r.text for r in paras[i + 1].runs)
        t2 = "".join(r.text for r in paras[i + 2].runs)
        t3 = "".join(r.text for r in paras[i + 3].runs)
        if (
            TARGET in t0
            and TITULO_TEMPLATE in t1
            and DESCR_TEMPLATE in t2
            and ENDFOR in t3
        ):
            _apply_titulo_bold_and_spacing(paras[i + 1])
            print("  asuntos_tratados: four-paragraph layout — bold + spacing on title")
            found = True
            break

    # Case B — legacy three paragraphs (single body line titulo: descr)
    if not found:
        paras = list(doc.paragraphs)
        for i in range(len(paras) - 2):
            t0 = "".join(r.text for r in paras[i].runs)
            t1 = "".join(r.text for r in paras[i + 1].runs)
            t2 = "".join(r.text for r in paras[i + 2].runs)
            if TARGET not in t0:
                continue
            if LEGACY_CONTENT not in t1 or ENDFOR not in t2:
                continue
            style_name = paras[i + 1].style.name
            parent = paras[i + 1]._element.getparent()
            mid_el = paras[i + 1]._element
            mid_idx = list(parent).index(mid_el)

            parent.remove(mid_el)
            _insert_jinja_paragraph(
                parent, mid_idx, style_name, TITULO_TEMPLATE, space_before_pt=12
            )
            _insert_jinja_paragraph(parent, mid_idx + 1, style_name, DESCR_TEMPLATE)

            _bind_titulo_paragraph()
            print(
                "  Upgraded 3-paragraph layout to 4 "
                "(split legacy titulo:descr → titulo bold + descripcion)"
            )
            found = True
            break

    # Case C — one paragraph contains both {%p for %} and {%p endfor %}
    if not found:
        for para in doc.paragraphs:
            text = "".join(r.text for r in para.runs)
            if "{%p for a in asuntos_tratados %}" not in text:
                continue
            if "{%p endfor %}" not in text:
                continue

            style_name = para.style.name
            parent = para._element.getparent()
            idx = list(parent).index(para._element)

            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = TARGET
            else:
                para.add_run(TARGET)

            _insert_jinja_paragraph(
                parent, idx + 1, style_name, TITULO_TEMPLATE, space_before_pt=12
            )
            _insert_jinja_paragraph(parent, idx + 2, style_name, DESCR_TEMPLATE)
            _insert_jinja_paragraph(parent, idx + 3, style_name, ENDFOR)

            _bind_titulo_paragraph()
            print(
                "  Split single paragraph into 4 "
                "(for / titulo bold + spacing / descripcion / endfor)"
            )
            found = True
            break

    if not found:
        print(
            "  asuntos_tratados: no matching paragraph pattern — skipping Step 3"
        )

    print("\n-- Step 4: Page margins (header/footer clearance) --")
    section = doc.sections[0]
    section.top_margin = Cm(4.5)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    print("  top=4.5cm, bottom=3.5cm, left=2.5cm, right=2.5cm (section 0)")

    doc.save(path)
    print(f"\nSaved: {path}")


if __name__ == "__main__":
    fix_document(TEMPLATE_PATH)
